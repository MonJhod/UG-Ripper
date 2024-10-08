"""
## UG-Ripper

# Author: MonJhod
# GitHub: https://github.com/MonJhod
# License: MIT
# Last Modified: 2024-08-19

`ug-ripper.py` is a web scraper designed to automate the process of downloading guitar 
tabs from Ultimate Guitar. It uses Selenium for web automation and BeautifulSoup for 
parsing HTML content. The program logs into Ultimate Guitar using user-provided credentials, 
navigates to a specified playlist, and downloads the guitar tabs as PDF files using pdfkit. 
The download location and other configurations are specified in a configuration file 
`config.ini`. The program handles login failures and timeouts gracefully, logging any 
errors encountered during the scraping process.

# Example Usage:
    python ug-ripper.py
"""

import configparser
import getpass
import logging
import os
import re
import time

import html2text
import pdfkit
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException, TimeoutException
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Read configuration from file
config = configparser.ConfigParser()
config.read('config.ini')

# Get download location from config
download_location = config.get('Download', 'location', fallback='.')

# Setup pdfkit config pointing to executable
pdfkitConfig = pdfkit.configuration(wkhtmltopdf=config.get('PDFKit', 'executable_path'))
# Setup pdfkit options to use UTF-8 encoding
pdfkitOptions = {
    'encoding': 'UTF-8',
    'page-size': 'Letter'
}

# List to keep track of failed downloads
failed_songs = []


def validate_config(config):
    """
    Validates the given configuration object to ensure all required sections and options are present.

    Args:
        config (ConfigParser): The configuration object to validate.

    Raises:
        ValueError: If a required section or option is missing from the configuration.
    """
    required_sections = {
        'Download': ['location'],
        'PDFKit': ['executable_path'],
        'URLs': ['login_url', 'playlist_url']
    }
    for section, options in required_sections.items():
        if not config.has_section(section):
            raise ValueError(f"Missing required section: {section}")
        for option in options:
            if not config.has_option(section, option):
                raise ValueError(f"Missing required option: {option} in section {section}")

def get_credentials():
    """
    Gets the user's login credentials for Ultimate Guitar.
    
    Returns:
        tuple: A tuple containing the username and password entered by the user.
    """
        
    username = input("Enter your username: ")
    password = getpass.getpass("Enter your password: ")
    return username, password

def preprocess_html(html):
    """
    Preprocess the HTML content by removing class attributes, ensuring <pre> tags are not encapsulated by <code> tags, and formatting the HTML.

    Args:
        html (str): The HTML content to prettify.

    Returns:
        str: The prettified HTML content.
    """
    soup = BeautifulSoup(html, 'html.parser')

    # Remove all class attributes
    for tag in soup.find_all(True):
        del tag['class']

    # Ensure <code> never encapsulates <pre> tags
    for code_tag in soup.find_all('code'):
        if code_tag.find('pre'):
            code_tag.unwrap()
            
    # Prettify the HTML content
    prettified_html = soup.prettify()
    
    return prettified_html

def setup_webdriver(download_location):
    """
    Sets up a Firefox WebDriver instance with the specified download location and other configuration options.
    
    Args:
        download_location (str): The path to the directory where downloaded files should be saved.
    
    Returns:
        webdriver.Firefox: A configured Firefox WebDriver instance.
    """
        
    profile = webdriver.FirefoxProfile()
    # Set the download folder list to 2, which means a custom location specified by the user
    profile.set_preference("browser.download.folderList", 2)
    # Disable showing the download manager when a download starts
    profile.set_preference("browser.download.manager.showWhenStarting", False)
    # Set the download directory to the absolute path of the specified download location
    profile.set_preference("browser.download.dir", os.path.abspath(download_location))
    # Automatically download files of type 'application/pdf' without asking
    profile.set_preference("browser.helperApps.neverAsk.saveToDisk", "application/pdf")
    
    options = webdriver.FirefoxOptions()
    options.profile = profile
    # Add the headless argument to run Firefox in headless mode (without a GUI)
    options.add_argument("--headless")
    # Disable automatic downloads
    options.enable_downloads = False
    
    driver = webdriver.Firefox(options=options)
    return driver

def login(username, password):
    """
    Logs in to the Ultimate Guitar website using the provided username and password.
    
    Args:
        username (str): The username to use for login.
        password (str): The password to use for login.
    
    Returns:
        webdriver.Firefox: The Firefox WebDriver instance if the login is successful, otherwise None.
    """
        
    login_url = config.get('URLs', 'login_url')
    driver = setup_webdriver(download_location)
    driver.get(login_url)
    
    try:
        login_button = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//button[span[text()='Log in']]")))
        login_button.click()
        
        username_field = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.NAME, 'username')))
        password_field = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.NAME, 'password')))
        
        username_field.send_keys(username)
        password_field.send_keys(password)
        password_field.send_keys(Keys.RETURN)
        
        # This is a cookie that is created after a successful login
        WebDriverWait(driver, 5).until(lambda d: d.get_cookie('bbusername') is not None)
        
        if driver.get_cookie('bbusername') is None:
            logging.error("Login failed.")
            driver.quit()
            return None
        else:
            logging.info("Login successful.")
            return driver
    except TimeoutException:
        logging.error("Timeout occurred while waiting for elements.")
        driver.quit()
        return None

def parse_song_links(driver):
    """
    Parses song links from the playlist URL using the provided WebDriver.

    Args:
        driver (selenium.webdriver.remote.webdriver.WebDriver): The WebDriver instance used to navigate and fetch the page source.

    Returns:
        list: A list of song URLs found on the playlist page.

    Raises:
        TimeoutException: If the page does not load within the specified timeout period.
    """
    playlist_url = config.get('URLs', 'playlist_url')
    driver.get(playlist_url)
    
    try:
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, 'a')))
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        song_links = soup.find_all('a', href=True)
        song_urls = [link['href'] for link in song_links if 'tabs.ultimate-guitar.com/tab' in link['href'] or 'https://tabs.ultimate-guitar.com/user/tab' in link['href']]
        return song_urls
    except TimeoutException:
        logging.error("Timeout occurred while waiting for song links.")
        return []

def dumb_docx_space_fixer(text):
    """
    Removes the extra newline and 4 spaces that follow each line.

    Args:
        text (str): The text to be fixed.
    """
    return re.sub(r'\r\n\s{4}', '\n', text)

def download_file(driver, song_url):
    """
    Downloads the song from the given URL and saves it as a PDF.

    Args:
        driver (selenium.webdriver): The Selenium WebDriver instance.
        song_url (str): The URL of the song to be downloaded.

    Returns:
        None

    Raises:
        TimeoutException: If the page takes too long to load.
        NoSuchElementException: If the required elements are not found on the page.
    """
    driver.get(song_url)
    
    try:
        # Wait for the page to load
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "/html/body/div[1]/div[3]/main/div[2]/article[1]/section[2]/article/div")))
        
        # Sleep for a short time since the full songs were not loading
        time.sleep(1)
        
        # Scroll to the bottom of the page to load the full content
        driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.END)
        
        # Extract song title
        song_title_element = driver.find_element(By.XPATH, "/html/body/div[1]/div[3]/main/div[2]/article[1]/section[1]/header")
        song_title = song_title_element.text.strip().replace('/', '_')    
        # Download everything the content xpaths 
        song_html = driver.find_element(By.XPATH, "/html/body/div[1]/div[3]/main/div[2]/article[1]/section[2]/article/div/section").get_attribute('innerHTML')
    
        # Concatenate song_title_element and song_html for the full download_html
        download_html = song_title_element.get_attribute('outerHTML') + song_html
        
        if config.getboolean("Download", "docx"):
            # Preprocess HTML by converting it to text with html2text
            text_maker = html2text.HTML2Text()
            text_maker.unicode_snob = True
            text_maker.ignore_links = True
            text_maker.single_line_break = True
            
            title_text = dumb_docx_space_fixer(text_maker.handle(preprocess_html(song_title_element.get_attribute('outerHTML'))))
            body_text = dumb_docx_space_fixer(text_maker.handle(preprocess_html(song_html)))
            doc = Document()
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            header_obj = doc.add_heading(level=2)
            header_obj.paragraph_format.left_indent = Inches(0)
            header_obj.paragraph_format.right_indent = Inches(0)
            header_obj.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            header = header_obj.add_run(title_text)
            header.font.name = "Times New Roman"
            header.font.color.rgb = RGBColor(0, 0, 0) # Black color
            body_obj = doc.add_paragraph()
            body_obj.paragraph_format.left_indent = Inches(0)
            body_obj.paragraph_format.right_indent = Inches(0)
            body_obj.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            body = body_obj.add_run(body_text)
            body.font.name = "Courier New"
            body.font.size = Pt(9.5)
            
            doc.save(f"{download_location}\\{song_title}.docx")
            
            logging.info(f"DOCX downloaded for song: {song_title}")
        else:             
            # Convert HTML to PDF
            pdfkit.from_string(download_html, f"{download_location}\\{song_title}.pdf", configuration=pdfkitConfig, options=pdfkitOptions)
            
            logging.info(f"PDF downloaded for song: {song_title}")
    except (TimeoutException, NoSuchElementException) as e:
        logging.error(f"Failed to download song: {driver.title.split(' | ')[0]}")
        logging.error(str(e))
        failed_songs.append(song_url)
        
def fetch_songs(driver):
    """
    Fetches and processes songs from the playlist URL using the provided WebDriver.

    Args:
        driver (selenium.webdriver.remote.webdriver.WebDriver): The WebDriver instance used to navigate and fetch the page source.

    Returns:
        None
    """
    song_urls = parse_song_links(driver)
    
    total_songs = len(song_urls)
    for index, song_url in enumerate(song_urls, start=1):
        logging.info(f"Processing song {index}/{total_songs}")
        download_file(driver, song_url)

if __name__ == "__main__":
    validate_config(config)
    while True:
        # Check if the config contains an Authentication section with Username and Password
        if config.has_section('Authentication') and config.has_option('Authentication', 'Username') and config.has_option('Authentication', 'Password'):
            username = config.get('Authentication', 'Username')
            password = config.get('Authentication', 'Password')
        else:
            username, password = get_credentials()
            
        driver = login(username, password)
        if driver:
            fetch_songs(driver)
            driver.quit()
            break
        else:
            logging.error("Login failed. Please try again.")
            
    if failed_songs:
        logging.info("Failed to download the following songs:")
        for song in failed_songs:
            logging.info(song)
    